from docx import Document
from docx.shared import Pt
from html import unescape

# Example text
text = "This is some <b>bold<b> and <i>italic<i> text."

# Create a new Word document
doc = Document()

# Add a paragraph with the formatted text
p = doc.add_paragraph()
for word in text.split():
    if word.startswith('<b>') and word.endswith('<b>'):
        # Add the bold formatting to the text
        p.add_run(unescape(word[3:-3])+" ").bold = True
    elif word.startswith('<i>') and word.endswith('<i>'):
        # Add the italic formatting to the text
        p.add_run(unescape(word[3:-3])+" ").italic = True
    else:
        # Add the text as is
        p.add_run(unescape(word)+" ")

# Set the font size of the bold text
for run in p.runs:
    if run.bold:
        run.font.size = Pt(14)

# Save the document
doc.save('formatted_text.docx')